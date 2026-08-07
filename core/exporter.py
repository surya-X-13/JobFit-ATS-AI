"""
core/exporter.py
ATS Document Exporter Engine.
Generates 100% ATS-compliant PDF (via ReportLab) and DOCX (via python-docx) documents.
"""
import io
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    HRFlowable,
    Table,
    TableStyle,
)
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY


def export_to_docx(resume_data: dict) -> bytes:
    """
    Generate an ATS-optimized editable DOCX document from structured resume data.
    """
    doc = docx.Document()

    # Set 0.75 inch margins for standard ATS readability
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Base Colors & Fonts
    PRIMARY_COLOR = RGBColor(14, 14, 36)      # Dark Navy/Slate
    ACCENT_COLOR = RGBColor(138, 124, 248)    # Indigo Accent
    TEXT_COLOR = RGBColor(40, 40, 40)         # Off-black text

    # Header Name
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = p_name.add_run(resume_data.get("full_name", "CANDIDATE NAME"))
    run_name.font.name = 'Arial'
    run_name.font.size = Pt(22)
    run_name.font.bold = True
    run_name.font.color.rgb = PRIMARY_COLOR

    # Contact Info
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_contact = p_contact.add_run(resume_data.get("contact_info", ""))
    run_contact.font.name = 'Arial'
    run_contact.font.size = Pt(9.5)
    run_contact.font.color.rgb = RGBColor(100, 100, 100)

    p_contact.paragraph_format.space_after = Pt(14)

    def add_section_header(title: str):
        p_head = doc.add_paragraph()
        run_head = p_head.add_run(title.upper())
        run_head.font.name = 'Arial'
        run_head.font.size = Pt(12)
        run_head.font.bold = True
        run_head.font.color.rgb = PRIMARY_COLOR
        p_head.paragraph_format.space_before = Pt(12)
        p_head.paragraph_format.space_after = Pt(4)

    # 1. Professional Summary
    if summary := resume_data.get("professional_summary"):
        add_section_header("PROFESSIONAL SUMMARY")
        p_sum = doc.add_paragraph()
        r_sum = p_sum.add_run(summary)
        r_sum.font.name = 'Arial'
        r_sum.font.size = Pt(10)
        r_sum.font.color.rgb = TEXT_COLOR
        p_sum.paragraph_format.space_after = Pt(10)

    # 2. Technical Skills
    if skills := resume_data.get("technical_skills"):
        add_section_header("TECHNICAL SKILLS")
        for category, skill_list in skills.items():
            p_sk = doc.add_paragraph()
            r_cat = p_sk.add_run(f"• {category}: ")
            r_cat.font.name = 'Arial'
            r_cat.font.size = Pt(10)
            r_cat.font.bold = True
            r_cat.font.color.rgb = PRIMARY_COLOR

            r_val = p_sk.add_run(", ".join(skill_list) if isinstance(skill_list, list) else str(skill_list))
            r_val.font.name = 'Arial'
            r_val.font.size = Pt(10)
            r_val.font.color.rgb = TEXT_COLOR
            p_sk.paragraph_format.space_after = Pt(3)

    # 3. Work Experience
    if exp_list := resume_data.get("work_experience"):
        add_section_header("WORK EXPERIENCE")
        for exp in exp_list:
            p_title = doc.add_paragraph()
            r_title = p_title.add_run(f"{exp.get('title', '')} | {exp.get('company', '')}")
            r_title.font.name = 'Arial'
            r_title.font.size = Pt(10.5)
            r_title.font.bold = True
            r_title.font.color.rgb = PRIMARY_COLOR

            if dates := exp.get("location_dates"):
                r_dates = p_title.add_run(f" ({dates})")
                r_dates.font.name = 'Arial'
                r_dates.font.size = Pt(9.5)
                r_dates.font.italic = True
                r_dates.font.color.rgb = RGBColor(100, 100, 100)

            p_title.paragraph_format.space_before = Pt(6)
            p_title.paragraph_format.space_after = Pt(2)

            for bullet in exp.get("bullet_points", []):
                p_bullet = doc.add_paragraph(style='List Bullet')
                r_bullet = p_bullet.add_run(bullet)
                r_bullet.font.name = 'Arial'
                r_bullet.font.size = Pt(9.5)
                r_bullet.font.color.rgb = TEXT_COLOR
                p_bullet.paragraph_format.space_after = Pt(2)

    # 4. Projects
    if proj_list := resume_data.get("projects"):
        add_section_header("KEY PROJECTS")
        for proj in proj_list:
            p_p = doc.add_paragraph()
            r_pn = p_p.add_run(f"• {proj.get('name', '')}: ")
            r_pn.font.name = 'Arial'
            r_pn.font.size = Pt(10)
            r_pn.font.bold = True
            r_pn.font.color.rgb = PRIMARY_COLOR

            r_pd = p_p.add_run(proj.get("details", ""))
            r_pd.font.name = 'Arial'
            r_pd.font.size = Pt(9.5)
            r_pd.font.color.rgb = TEXT_COLOR
            p_p.paragraph_format.space_after = Pt(3)

    # 5. Education
    if edu_list := resume_data.get("education"):
        add_section_header("EDUCATION")
        for edu in edu_list:
            p_ed = doc.add_paragraph()
            r_deg = p_ed.add_run(f"{edu.get('degree', '')}")
            r_deg.font.name = 'Arial'
            r_deg.font.size = Pt(10)
            r_deg.font.bold = True
            r_deg.font.color.rgb = PRIMARY_COLOR

            if inst := edu.get("institution_dates"):
                r_inst = p_ed.add_run(f" - {inst}")
                r_inst.font.name = 'Arial'
                r_inst.font.size = Pt(9.5)
                r_inst.font.color.rgb = TEXT_COLOR
            p_ed.paragraph_format.space_after = Pt(3)

    # Save to Bytes buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def export_to_pdf(resume_data: dict) -> bytes:
    """
    Generate an ATS-optimized 1-column PDF document using ReportLab.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        leftMargin=54,  # 0.75 in
        rightMargin=54,
        topMargin=54,
        bottomMargin=54,
    )

    styles = getSampleStyleSheet()

    # Custom ReportLab Paragraph Styles
    PRIMARY_PDF = colors.HexColor("#0e0e24")
    TEXT_PDF = colors.HexColor("#282828")
    MUTED_PDF = colors.HexColor("#646464")

    style_name = ParagraphStyle(
        'DocName',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=PRIMARY_PDF,
        alignment=TA_CENTER,
        spaceAfter=4,
    )

    style_contact = ParagraphStyle(
        'DocContact',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=12,
        textColor=MUTED_PDF,
        alignment=TA_CENTER,
        spaceAfter=12,
    )

    style_section = ParagraphStyle(
        'DocSection',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=14,
        textColor=PRIMARY_PDF,
        spaceBefore=10,
        spaceAfter=4,
    )

    style_body = ParagraphStyle(
        'DocBody',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9.5,
        leading=13.5,
        textColor=TEXT_PDF,
        spaceAfter=6,
    )

    style_bullet = ParagraphStyle(
        'DocBullet',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=9,
        leading=13,
        textColor=TEXT_PDF,
        leftIndent=14,
        firstLineIndent=-10,
        spaceAfter=3,
    )

    story = []

    # Candidate Name & Contact
    story.append(Paragraph(resume_data.get("full_name", "CANDIDATE NAME"), style_name))
    story.append(Paragraph(resume_data.get("contact_info", ""), style_contact))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor("#8a7cf8"), spaceAfter=10))

    # 1. Professional Summary
    if summary := resume_data.get("professional_summary"):
        story.append(Paragraph("PROFESSIONAL SUMMARY", style_section))
        story.append(Paragraph(summary, style_body))

    # 2. Technical Skills
    if skills := resume_data.get("technical_skills"):
        story.append(Paragraph("TECHNICAL SKILLS", style_section))
        for cat, sk_list in skills.items():
            s_str = ", ".join(sk_list) if isinstance(sk_list, list) else str(sk_list)
            text = f"<b>{cat}:</b> {s_str}"
            story.append(Paragraph(text, style_body))

    # 3. Work Experience
    if exp_list := resume_data.get("work_experience"):
        story.append(Paragraph("WORK EXPERIENCE", style_section))
        for exp in exp_list:
            title_comp = f"<b>{exp.get('title','')}</b> | {exp.get('company','')}"
            if dates := exp.get("location_dates"):
                title_comp += f" <i>({dates})</i>"
            story.append(Paragraph(title_comp, style_body))

            for bullet in exp.get("bullet_points", []):
                story.append(Paragraph(f"• {bullet}", style_bullet))
            story.append(Spacer(1, 4))

    # 4. Key Projects
    if proj_list := resume_data.get("projects"):
        story.append(Paragraph("KEY PROJECTS", style_section))
        for proj in proj_list:
            p_text = f"<b>• {proj.get('name','')}:</b> {proj.get('details','')}"
            story.append(Paragraph(p_text, style_body))

    # 5. Education
    if edu_list := resume_data.get("education"):
        story.append(Paragraph("EDUCATION", style_section))
        for edu in edu_list:
            e_text = f"<b>{edu.get('degree','')}</b> - {edu.get('institution_dates','')}"
            story.append(Paragraph(e_text, style_body))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()
